import os
import numpy as np
from math import ceil
import pandas as pd
from collections import defaultdict

import string

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


def get_column_name(columns, options):
	return set(columns).intersection(set(options)).pop()


WCI = {-1: WD_COLOR_INDEX.RED,
		0: WD_COLOR_INDEX.YELLOW, \
		1: WD_COLOR_INDEX.BRIGHT_GREEN, \
		2: WD_COLOR_INDEX.GRAY_25, \
		3: WD_COLOR_INDEX.TEAL}

def set_global_voc_cost(c):
	global GOLBAL_VOC_COST
	GOLBAL_VOC_COST = c

def log_star(x):
	"""
	Universal code length

	"""
	return 2 * ceil(np.log2(x)) + 1 if x != 0 else 0

def word_cost():
	return GOLBAL_VOC_COST

def sequence_cost(seq):
	"""
	Output encoding cost for a given sequence

	"""
	return log_star(len(seq)) + len(seq) * word_cost()

def str_prep(s):
	s = s.translate(str.maketrans('', '', string.punctuation)).split(' ')
	s = np.array([ss.lower() for ss in s if len(ss) != 0])
	return s

class Sequence:
	def __init__(self, sid, data, timetick):
		self.sid = sid
		self.data = data
		self.timetick = timetick
		self.cost = 0

	def update_cost(self):
		self.cost = sequence_cost(self.data)

def read_data(path):
	# df = pd.read_csv(path)
	df = pd.read_csv(path).iloc[:-1014]
	lsh_label = df['LSH label'].unique()
	sequences = {0: defaultdict(dict)}

	voc = set()
	length = []
	for label in lsh_label:
		# for id, text in df[df['LSH label'] == label][['id', 'text']].values:
		# for id, text in df[df['LSH label'] == label][['ad_id', 'description']].values:
		for id, text in df[df['LSH label'] == label][['ad_id', 'body']].values:
			try:
				text = str_prep(text)
				for t in text:
					voc.add(t)
			except:
				continue
			if len(text) != 0:
				sequences[0][0][id] = Sequence(id, text, 0)
				# sequences[0][label][id] = Sequence(id, text, 0)
				length.append(len(text))

	gvc = ceil(np.log2(len(voc)))
	set_global_voc_cost(gvc)
	for k1 in sequences[0].keys():
		for k2 in sequences[0][k1].keys():
			sequences[0][k1][k2].update_cost()

	return sequences

def read_temporal_data(path):
    df = pd.read_csv(path)
    batch_num = df['batch_num'].unique()
    sequences = defaultdict(dict)

    id_col = get_column_name(df.columns, ['id', 'ad_id'])
    text_col = get_column_name(df.columns, ['body', 'text', 'description'])

    voc = set()
    for bn in batch_num:
        tdf = df[df['batch_num'] == bn]
        lsh_label = tdf['LSH label'].unique()
        sequences[bn] = {}
        for label in lsh_label:
            sequences[bn][label] = {}
            for id, text in tdf[tdf['LSH label'] == label][[id_col, text_col]].values:
                try:
                    text = str_prep(text)
                    for t in text:
                        voc.add(t)
                except:
                    continue
                if len(text) != 0:
                    sequences[bn][label][id] = Sequence(id, text, bn)

    gvc = ceil(np.log2(len(voc)))
    set_global_voc_cost(gvc)
    for k1 in sequences.keys():
        for k2 in sequences[k1].keys():
            for k3 in sequences[k1][k2].keys():
                sequences[k1][k2][k3].update_cost()

    return sequences

def read_test_data(path):
	df = pd.read_csv(path)
	sequences = defaultdict(dict)

	voc = set()
	for idx, (id, text) in enumerate(df[['ad_id', 'body']].values):
		try:
			text = str_prep(text)
			for t in text:
					voc.add(t)
		except:
			continue
		if len(text) != 0:
			try:
				sequences[int(idx / 2000)][0][id] = Sequence(id, text, int(idx / 2000))
			except:
				sequences[int(idx / 2000)] = {}
				sequences[int(idx / 2000)][0] = {}
				sequences[int(idx / 2000)][0][id] = Sequence(id, text, int(idx / 2000))

	gvc = ceil(np.log2(len(voc)))
	set_global_voc_cost(gvc)
	for k1 in sequences.keys():
		for k2 in sequences[k1].keys():
			for k3 in sequences[k1][k2].keys():
				sequences[k1][k2][k3].update_cost()

	return sequences

def read_test_data_old(path):
	df = pd.read_csv(path)
	sequences = defaultdict(dict)

	voc = set()
	df = df[df['LSH label'] == 1262]

	for idx, (id, text) in enumerate(df[['ad_id', 'body']].values):
		try:
			text = str_prep(text)
			for t in text:
					voc.add(t)
		except:
			continue
		if len(text) != 0:
			try:
				sequences[int(idx / 20)][0][id] = Sequence(id, text, int(idx / 20))
			except:
				sequences[int(idx / 20)] = {}
				sequences[int(idx / 20)][0] = {}
				sequences[int(idx / 20)][0][id] = Sequence(id, text, int(idx / 20))

	gvc = ceil(np.log2(len(voc)))
	set_global_voc_cost(gvc)
	for k1 in sequences.keys():
		for k2 in sequences[k1].keys():
			for k3 in sequences[k1][k2].keys():
				sequences[k1][k2][k3].update_cost()

	return sequences

def output_word(temp, word_path):
	"""
	Output highlight content with office word document

	"""
	### Initialize document
	doc = Document()
	proc = doc.add_paragraph()
	for s, c in zip(['Slot', 'Matched', 'Substitution', 'Deletion', 'Insertion'], WCI.values()):
		font = proc.add_run(s).font
		font.highlight_color = c
		proc.add_run(' ')

	### Template content
	proc = doc.add_paragraph()
	proc.add_run('Template: \n')
	proc.add_run(temp.template.seq())
	proc.add_run('\n\n-----------------------------------------------------------------\n')

	### Iterate all aligned sequences
	for cs in temp.conds.values():
		proc = doc.add_paragraph()
		for c, s in cs:
			font = proc.add_run(s).font
			font.highlight_color = WCI[c]
			proc.add_run(' ')

	doc.save(word_path)

def output_results(temp_arr, output_path, html_name='graph.html', word_name='text.docx'):
	"""
	Output template results

	"""
	if not os.path.exists(output_path):
		os.makedirs(output_path)

	### Iterate all templates
	for idx, temp in enumerate(temp_arr):
		temp_path = os.path.join(output_path, 'template_' + str(idx + 1) + '_timetick_' + str(list(temp.timetick)))
		if not os.path.exists(temp_path):
			os.makedirs(temp_path)

		### Output html
		temp.template.htmlOutput(open(os.path.join(temp_path, html_name), 'w'))

		### Output word document
		output_word(temp, os.path.join(temp_path, word_name))
